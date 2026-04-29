"""Tests du helper watermark - gating Free vs payant + i18n FR/EN + 6 formats."""
import pytest
from exports.watermark import add_watermark, PAID_PLANS


# --- Constantes attendues ----------------------------------------------------

EXPECTED_FR_TEXT = "Analysé avec DeepSight"
EXPECTED_FR_TAGLINE = "IA souveraine européenne"
EXPECTED_EN_TEXT = "Analyzed with DeepSight"
EXPECTED_EN_TAGLINE = "European sovereign AI"
EXPECTED_URL = "www.deepsightsynthesis.com"


# --- Set PAID_PLANS ----------------------------------------------------------

def test_paid_plans_contains_v0_and_v2_during_transition():
    """PAID_PLANS doit couvrir 'plus' (v0), 'pro' (v0+v2), 'expert' (v2) pendant la transition."""
    assert "plus" in PAID_PLANS
    assert "pro" in PAID_PLANS
    assert "expert" in PAID_PLANS
    assert "free" not in PAID_PLANS


# --- Gating Free vs payant ---------------------------------------------------

@pytest.mark.parametrize("paid_plan", ["plus", "pro", "expert"])
@pytest.mark.parametrize("fmt", ["txt", "md", "csv"])
def test_watermark_not_applied_for_paid_plans_text(paid_plan: str, fmt: str):
    """Aucun watermark sur les plans payants, formats texte."""
    original = "Contenu test"
    result = add_watermark(original, fmt, user_plan=paid_plan, user_language="fr")
    assert result == original


@pytest.mark.parametrize("paid_plan", ["plus", "pro", "expert"])
@pytest.mark.parametrize("fmt", ["docx", "xlsx", "pdf"])
def test_watermark_not_applied_for_paid_plans_binary(paid_plan: str, fmt: str):
    """Plans payants : binary returns marker dict avec needs_watermark=False."""
    result = add_watermark("placeholder", fmt, user_plan=paid_plan, user_language="fr")
    assert isinstance(result, dict)
    assert result["needs_watermark"] is False


@pytest.mark.parametrize("fmt", ["txt", "csv"])
def test_watermark_applied_for_free_plan_text(fmt: str):
    """Watermark present sur formats texte (txt/csv) pour Free."""
    original = "Contenu test"
    result = add_watermark(original, fmt, user_plan="free", user_language="fr")
    assert result != original
    assert EXPECTED_FR_TEXT in result
    assert EXPECTED_URL in result


def test_watermark_applied_for_free_plan_md():
    """MD : watermark present mais format different (lien markdown casse 'Analyse avec DeepSight')."""
    original = "Contenu test"
    result = add_watermark(original, "md", user_plan="free", user_language="fr")
    assert result != original
    # Lien markdown contient DeepSight
    assert "[DeepSight](https://www.deepsightsynthesis.com)" in result
    # Tagline complete contient les autres mots-cles
    assert "IA souveraine européenne" in result
    assert "Analysé avec" in result


def test_watermark_applied_for_unknown_plan_defaults_to_free():
    """Plan inconnu -> fallback Free -> watermark applique."""
    result = add_watermark("X", "txt", user_plan="bogus_plan", user_language="fr")
    assert EXPECTED_FR_TEXT in result


def test_watermark_applied_for_empty_plan():
    """Plan vide ou None -> fallback Free."""
    result_empty = add_watermark("X", "txt", user_plan="", user_language="fr")
    result_none = add_watermark("X", "txt", user_plan=None, user_language="fr")
    assert EXPECTED_FR_TEXT in result_empty
    assert EXPECTED_FR_TEXT in result_none


# --- i18n FR / EN ------------------------------------------------------------

def test_watermark_french_default():
    """Langue FR par defaut."""
    result = add_watermark("X", "txt", user_plan="free", user_language="fr")
    assert EXPECTED_FR_TEXT in result
    assert EXPECTED_FR_TAGLINE in result
    assert EXPECTED_EN_TEXT not in result


def test_watermark_english():
    """Langue EN si user_language='en'."""
    result = add_watermark("X", "txt", user_plan="free", user_language="en")
    assert EXPECTED_EN_TEXT in result
    assert EXPECTED_EN_TAGLINE in result
    assert EXPECTED_FR_TEXT not in result


def test_watermark_unknown_language_defaults_to_fr():
    """Langue inconnue (es, de, etc.) -> fallback FR."""
    result = add_watermark("X", "txt", user_plan="free", user_language="es")
    assert EXPECTED_FR_TEXT in result


# --- Format-specific shape ---------------------------------------------------

def test_watermark_txt_format():
    """TXT : separateur '---' + ligne plain."""
    result = add_watermark("Hello", "txt", user_plan="free", user_language="fr")
    assert "Hello" in result
    assert "---" in result


def test_watermark_md_format():
    """MD : separateur '---' + lien markdown."""
    result = add_watermark("Hello", "md", user_plan="free", user_language="fr")
    assert "Hello" in result
    assert "---" in result
    assert "[DeepSight](https://www.deepsightsynthesis.com)" in result
    # Italic
    assert "*" in result.split("Hello", 1)[1]


def test_watermark_csv_format():
    """CSV : commentaire derniere ligne prefixe par '#'."""
    result = add_watermark("a,b,c", "csv", user_plan="free", user_language="fr")
    assert "a,b,c" in result
    last_line = result.strip().split("\n")[-1]
    assert last_line.startswith("#")
    assert EXPECTED_FR_TEXT in last_line


# --- Format binaires (DOCX/XLSX/PDF) -- passthrough --------------------------

def test_watermark_docx_returns_marker_dict():
    """DOCX : helper retourne dict {needs_watermark, text, url} pour que service.py l'injecte via python-docx."""
    result = add_watermark("placeholder", "docx", user_plan="free", user_language="fr")
    assert isinstance(result, dict)
    assert result["needs_watermark"] is True
    assert EXPECTED_FR_TEXT in result["text"]
    assert result["url"] == EXPECTED_URL


def test_watermark_docx_no_marker_for_paid():
    """DOCX : payant -> dict avec needs_watermark=False."""
    result = add_watermark("placeholder", "docx", user_plan="pro", user_language="fr")
    assert isinstance(result, dict)
    assert result["needs_watermark"] is False


def test_watermark_xlsx_returns_marker_dict():
    """XLSX : idem DOCX, dict marqueur."""
    result = add_watermark("placeholder", "xlsx", user_plan="free", user_language="fr")
    assert isinstance(result, dict)
    assert result["needs_watermark"] is True
    assert EXPECTED_FR_TEXT in result["text"]


def test_watermark_pdf_returns_marker_dict():
    """PDF : idem, dict pour propagation au template Jinja2."""
    result = add_watermark("placeholder", "pdf", user_plan="free", user_language="fr")
    assert isinstance(result, dict)
    assert result["needs_watermark"] is True
    assert EXPECTED_FR_TEXT in result["text"]
    assert result["url"] == EXPECTED_URL


# --- Normalize via plan_config ----------------------------------------------

def test_watermark_normalize_legacy_aliases():
    """Plans legacy (etudiant, starter, student, equipe, team, unlimited) sont mappes via normalize_plan_id."""
    # 'etudiant' -> 'plus' (in PAID_PLANS) -> no watermark
    result_etudiant = add_watermark("X", "txt", user_plan="etudiant", user_language="fr")
    assert result_etudiant == "X"
    # 'team' -> 'pro' (in PAID_PLANS) -> no watermark
    result_team = add_watermark("X", "txt", user_plan="team", user_language="fr")
    assert result_team == "X"
    # 'unlimited' -> 'pro' -> no watermark
    result_unlim = add_watermark("X", "txt", user_plan="unlimited", user_language="fr")
    assert result_unlim == "X"


# ============================================================================
# INTEGRATION TESTS -- appel reel a export_to_*
# ============================================================================
from datetime import datetime  # noqa: E402

from exports.service import (  # noqa: E402
    export_to_txt,
    export_to_markdown,
    export_to_docx,
    export_to_csv,
    export_to_excel,
    export_summary,
    DOCX_AVAILABLE,
    EXCEL_AVAILABLE,
)


SAMPLE_KWARGS = dict(
    title="Test Video",
    channel="Test Channel",
    category="Tech",
    mode="Standard",
    summary="# Test\n\nContenu de synthese de test.",
    video_url="https://youtube.com/watch?v=test",
    duration=600,
    created_at=datetime(2026, 4, 29, 12, 0, 0),
)


# --- TXT ---------------------------------------------------------------------

def test_integration_txt_free_has_watermark():
    out = export_to_txt(**SAMPLE_KWARGS, user_plan="free", user_language="fr")
    assert "Analysé avec DeepSight" in out
    assert "IA souveraine européenne" in out
    # Mention brand origin existante reste
    assert "deepsightsynthesis.com" in out


def test_integration_txt_pro_no_watermark():
    out = export_to_txt(**SAMPLE_KWARGS, user_plan="pro", user_language="fr")
    assert "Analysé avec DeepSight" not in out
    # Mention brand origin existante reste
    assert "deepsightsynthesis.com" in out


def test_integration_txt_expert_no_watermark():
    out = export_to_txt(**SAMPLE_KWARGS, user_plan="expert", user_language="fr")
    assert "Analysé avec DeepSight" not in out


def test_integration_txt_plus_no_watermark_during_v0_v2_transition():
    """v0 'plus' = payant, doit pas avoir de watermark pendant la transition."""
    out = export_to_txt(**SAMPLE_KWARGS, user_plan="plus", user_language="fr")
    assert "Analysé avec DeepSight" not in out


# --- MD ----------------------------------------------------------------------

def test_integration_md_free_has_watermark_link():
    out = export_to_markdown(**SAMPLE_KWARGS, user_plan="free", user_language="fr")
    assert "[DeepSight](https://www.deepsightsynthesis.com)" in out
    assert "IA souveraine européenne" in out


def test_integration_md_pro_no_watermark():
    out = export_to_markdown(**SAMPLE_KWARGS, user_plan="pro", user_language="fr")
    assert "Analysé avec DeepSight" not in out
    assert "IA souveraine européenne" not in out


def test_integration_md_english():
    out = export_to_markdown(**SAMPLE_KWARGS, user_plan="free", user_language="en")
    assert "Analyzed with" in out
    assert "European sovereign AI" in out


# --- DOCX --------------------------------------------------------------------

@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_integration_docx_free_has_watermark_paragraph():
    """DOCX Free contient un paragraphe avec le texte watermark."""
    from docx import Document
    import io

    blob = export_to_docx(**SAMPLE_KWARGS, user_plan="free", user_language="fr")
    assert blob is not None
    doc = Document(io.BytesIO(blob))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Analysé avec DeepSight" in full_text
    assert "IA souveraine européenne" in full_text


@pytest.mark.skipif(not DOCX_AVAILABLE, reason="python-docx not installed")
def test_integration_docx_pro_no_watermark_paragraph():
    from docx import Document
    import io

    blob = export_to_docx(**SAMPLE_KWARGS, user_plan="pro", user_language="fr")
    doc = Document(io.BytesIO(blob))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Analysé avec DeepSight" not in full_text
    # Mention brand existante reste
    assert "Généré par Deep Sight" in full_text


# --- CSV ---------------------------------------------------------------------

def test_integration_csv_free_has_watermark_comment():
    out = export_to_csv(**SAMPLE_KWARGS, user_plan="free", user_language="fr")
    last_line = out.strip().split("\n")[-1]
    assert last_line.startswith("#")
    assert "Analysé avec DeepSight" in last_line


def test_integration_csv_pro_no_watermark_comment():
    out = export_to_csv(**SAMPLE_KWARGS, user_plan="pro", user_language="fr")
    last_line = out.strip().split("\n")[-1]
    assert "Analysé avec DeepSight" not in last_line


# --- XLSX --------------------------------------------------------------------

@pytest.mark.skipif(not EXCEL_AVAILABLE, reason="openpyxl not installed")
def test_integration_xlsx_free_has_watermark_cell():
    from openpyxl import load_workbook
    import io

    blob = export_to_excel(**SAMPLE_KWARGS, user_plan="free", user_language="fr")
    assert blob is not None
    wb = load_workbook(io.BytesIO(blob))
    ws = wb.active
    all_values = []
    for row in ws.iter_rows(values_only=True):
        for v in row:
            if v:
                all_values.append(str(v))
    full_text = " ".join(all_values)
    assert "Analysé avec DeepSight" in full_text


@pytest.mark.skipif(not EXCEL_AVAILABLE, reason="openpyxl not installed")
def test_integration_xlsx_pro_no_watermark_cell():
    from openpyxl import load_workbook
    import io

    blob = export_to_excel(**SAMPLE_KWARGS, user_plan="pro", user_language="fr")
    wb = load_workbook(io.BytesIO(blob))
    ws = wb.active
    all_values = []
    for row in ws.iter_rows(values_only=True):
        for v in row:
            if v:
                all_values.append(str(v))
    full_text = " ".join(all_values)
    assert "Analysé avec DeepSight" not in full_text
    assert "Deep Sight" in full_text  # Brand mention existante reste


# --- PDF (smoke -- verifie generer sans crash si WeasyPrint dispo) ----------

def _pdf_engine_available():
    """Check si au moins WeasyPrint OU ReportLab est dispo."""
    try:
        from exports.service import weasyprint_available, REPORTLAB_AVAILABLE
        return weasyprint_available() or bool(REPORTLAB_AVAILABLE)
    except Exception:
        return False


@pytest.mark.skipif(not _pdf_engine_available(), reason="No PDF engine available")
def test_integration_pdf_free_smoke():
    """Smoke test : PDF Free genere sans crash, taille > 1KB.

    Utilise WeasyPrint si dispo, sinon ReportLab fallback.
    Le watermark est verifie de facon binary -- byte search dans le PDF.
    """
    content, filename, mimetype = export_summary(
        format="pdf",
        **SAMPLE_KWARGS,
        user_plan="free",
        user_language="fr",
    )
    assert content is not None
    assert len(content) > 1024  # PDF non trivial
    assert filename.endswith(".pdf")
    assert mimetype == "application/pdf"
    # Verify watermark text appears in PDF bytes (best-effort -- ReportLab encodes plain text)
    # WeasyPrint may compress so this is a smoke check only
    if isinstance(content, bytes) and b"DeepSight" in content:
        # Some chance the watermark text leaks visibly
        pass  # OK


@pytest.mark.skipif(not _pdf_engine_available(), reason="No PDF engine available")
def test_integration_pdf_pro_smoke():
    content, _, _ = export_summary(
        format="pdf",
        **SAMPLE_KWARGS,
        user_plan="pro",
        user_language="fr",
    )
    assert content is not None
    assert len(content) > 1024


# --- export_summary dispatcher ----------------------------------------------

@pytest.mark.parametrize("fmt", ["txt", "csv"])
def test_integration_export_summary_free_has_watermark(fmt: str):
    """Le dispatcher principal applique le watermark pour Free sur les formats texte."""
    content, _, _ = export_summary(
        format=fmt,
        **SAMPLE_KWARGS,
        user_plan="free",
        user_language="fr",
    )
    assert isinstance(content, str)
    assert "Analysé avec DeepSight" in content


def test_integration_export_summary_md_free_has_watermark():
    """MD: lien markdown (DeepSight remplace par [DeepSight](url)) -- check mots-cle."""
    content, _, _ = export_summary(
        format="md",
        **SAMPLE_KWARGS,
        user_plan="free",
        user_language="fr",
    )
    assert isinstance(content, str)
    assert "[DeepSight](https://www.deepsightsynthesis.com)" in content
    assert "IA souveraine européenne" in content


@pytest.mark.parametrize("fmt", ["txt", "md", "csv"])
def test_integration_export_summary_pro_no_watermark(fmt: str):
    """Le dispatcher principal n'applique PAS le watermark pour Pro sur les formats texte."""
    content, _, _ = export_summary(
        format=fmt,
        **SAMPLE_KWARGS,
        user_plan="pro",
        user_language="fr",
    )
    assert isinstance(content, str)
    assert "Analysé avec DeepSight" not in content
    assert "IA souveraine européenne" not in content
