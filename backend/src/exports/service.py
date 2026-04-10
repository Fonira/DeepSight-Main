"""
╔════════════════════════════════════════════════════════════════════════════════════╗
║  📄 EXPORT SERVICE — Génération PDF, DOCX, TXT, Markdown, CSV, Excel               ║
╠════════════════════════════════════════════════════════════════════════════════════╣
║  v2.0 — Professional PDF exports with branded design + CSV/Excel                   ║
║  • WeasyPrint for beautiful HTML→PDF rendering                                     ║
║  • Multiple export modes (full, summary, flashcards, study pack)                   ║
║  • Fallback to ReportLab if WeasyPrint unavailable                                 ║
║  • CSV + Excel exports for data analysis                                           ║
╚════════════════════════════════════════════════════════════════════════════════════╝
"""

import os
import io
import re
import csv
from datetime import datetime
from typing import Optional, Dict, Any, List, Tuple
from pathlib import Path

from core.http_client import shared_http_client

# ═══════════════════════════════════════════════════════════════════════════════
# 📦 IMPORTS LAZY — Chargés uniquement quand nécessaire (économie ~80MB RAM)
# Railway 512MB : chaque MB compte, ces libs ne sont utilisées que pour l'export
# ═══════════════════════════════════════════════════════════════════════════════

# Flags d'état — initialisés à None (pas encore testé)
_PDF_CHECKED = False
_DOCX_CHECKED = False
_REPORTLAB_CHECKED = False
_EXCEL_CHECKED = False

# Modules lazy (None = pas encore chargé)
_pdf_module = None
_docx_module = None
_reportlab_modules = None
_excel_module = None


def _ensure_pdf():
    """Lazy load du PDF generator (WeasyPrint)"""
    global _PDF_CHECKED, _pdf_module
    if _PDF_CHECKED:
        return _pdf_module is not None
    _PDF_CHECKED = True
    try:
        from . import pdf_generator
        _pdf_module = pdf_generator
        return True
    except (ImportError, OSError):
        return False


def _ensure_docx():
    """Lazy load de python-docx"""
    global _DOCX_CHECKED, _docx_module
    if _DOCX_CHECKED:
        return _docx_module is not None
    _DOCX_CHECKED = True
    try:
        import docx as _d
        _docx_module = _d
        return True
    except ImportError:
        return False


def _ensure_reportlab():
    """Lazy load de ReportLab"""
    global _REPORTLAB_CHECKED, _reportlab_modules
    if _REPORTLAB_CHECKED:
        return _reportlab_modules is not None
    _REPORTLAB_CHECKED = True
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.units import cm
        _reportlab_modules = {
            "A4": A4, "getSampleStyleSheet": getSampleStyleSheet,
            "ParagraphStyle": ParagraphStyle, "HexColor": HexColor,
            "SimpleDocTemplate": SimpleDocTemplate, "Paragraph": Paragraph,
            "Spacer": Spacer, "PageBreak": PageBreak, "cm": cm,
        }
        return True
    except ImportError:
        return False


def _ensure_excel():
    """Lazy load de openpyxl"""
    global _EXCEL_CHECKED, _excel_module
    if _EXCEL_CHECKED:
        return _excel_module is not None
    _EXCEL_CHECKED = True
    try:
        import openpyxl as _xl
        _excel_module = _xl
        return True
    except ImportError:
        return False


# ── Fonctions de compatibilité (remplacent les anciens flags booléens) ──

def weasyprint_available() -> bool:
    """Check si WeasyPrint est disponible (lazy load)"""
    return _ensure_pdf() and _pdf_module.is_pdf_available()

def generate_pdf_weasyprint(*args, **kwargs):
    """Proxy vers pdf_generator.generate_pdf (lazy)"""
    if not _ensure_pdf():
        raise ImportError("WeasyPrint non disponible")
    return _pdf_module.generate_pdf(*args, **kwargs)

def get_pdf_export_type():
    """Accès lazy à PDFExportType enum"""
    if _ensure_pdf():
        return _pdf_module.PDFExportType
    return None


class _LazyFlag:
    """Descriptor pour simuler un booléen avec lazy loading"""
    def __init__(self, checker):
        self._checker = checker
    def __bool__(self):
        return self._checker()
    def __repr__(self):
        return str(bool(self))

DOCX_AVAILABLE = _LazyFlag(_ensure_docx)
REPORTLAB_AVAILABLE = _LazyFlag(_ensure_reportlab)
EXCEL_AVAILABLE = _LazyFlag(_ensure_excel)


# ═══════════════════════════════════════════════════════════════════════════════
# 🎨 STYLES & CONFIGURATION
# ═══════════════════════════════════════════════════════════════════════════════

# Couleurs Deep Sight (thème océan/steampunk)
COLORS = {
    "primary": "#0D4F4F",      # Teal profond
    "secondary": "#D4A574",    # Cuivre/laiton
    "accent": "#00CED1",       # Cyan
    "text": "#1A1A2E",         # Texte sombre
    "light": "#F5F5F5",        # Fond clair
}

# Template de header pour les exports
HEADER_TEMPLATE = """
╔═══════════════════════════════════════════════════════════════════════════╗
║  🤿 DEEP SIGHT — Analyse Intelligente                                      ║
╚═══════════════════════════════════════════════════════════════════════════╝
"""


# ═══════════════════════════════════════════════════════════════════════════════
# 🔧 HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def format_duration(duration: int) -> str:
    """Formate une durée en secondes en string lisible"""
    if not duration:
        return "N/A"
    hours, remainder = divmod(duration, 3600)
    minutes, seconds = divmod(remainder, 60)
    if hours:
        return f"{hours}h{minutes:02d}m{seconds:02d}s"
    return f"{minutes}m{seconds:02d}s"


def clean_filename(title: str, timestamp: str) -> str:
    """Génère un nom de fichier sûr"""
    safe_title = re.sub(r'[^\w\s-]', '', title)[:50].strip()
    safe_title = re.sub(r'[-\s]+', '_', safe_title)
    return f"deepsight_{safe_title}_{timestamp}"


# ═══════════════════════════════════════════════════════════════════════════════
# 📝 EXPORT TXT
# ═══════════════════════════════════════════════════════════════════════════════

def export_to_txt(
    title: str,
    channel: str,
    category: str,
    mode: str,
    summary: str,
    video_url: str = "",
    duration: int = 0,
    created_at: datetime = None
) -> str:
    """Exporte l'analyse en format texte brut"""
    
    duration_str = format_duration(duration)
    date_str = created_at.strftime("%d/%m/%Y %H:%M") if created_at else datetime.now().strftime("%d/%m/%Y %H:%M")
    
    content = f"""{HEADER_TEMPLATE}
═══════════════════════════════════════════════════════════════════════════
📺 VIDÉO ANALYSÉE
═══════════════════════════════════════════════════════════════════════════

Titre    : {title}
Chaîne   : {channel}
Durée    : {duration_str}
Catégorie: {category}
Mode     : {mode}
URL      : {video_url}
Analysé  : {date_str}

═══════════════════════════════════════════════════════════════════════════
📋 SYNTHÈSE
═══════════════════════════════════════════════════════════════════════════

{summary}

═══════════════════════════════════════════════════════════════════════════
                    Généré par Deep Sight — deepsightsynthesis.com
═══════════════════════════════════════════════════════════════════════════
"""
    return content


# ═══════════════════════════════════════════════════════════════════════════════
# 📝 EXPORT MARKDOWN
# ═══════════════════════════════════════════════════════════════════════════════

def export_to_markdown(
    title: str,
    channel: str,
    category: str,
    mode: str,
    summary: str,
    video_url: str = "",
    duration: int = 0,
    thumbnail_url: str = "",
    entities: Dict = None,
    reliability_score: float = None,
    created_at: datetime = None,
    flashcards: List[Dict] = None
) -> str:
    """Exporte l'analyse en format Markdown"""
    
    duration_str = format_duration(duration) if duration else "N/A"
    date_str = created_at.strftime("%d/%m/%Y à %H:%M") if created_at else datetime.now().strftime("%d/%m/%Y à %H:%M")
    
    content = f"""# 🤿 Deep Sight — Analyse

---

## 📺 Vidéo analysée

| Propriété | Valeur |
|-----------|--------|
| **Titre** | {title} |
| **Chaîne** | {channel} |
| **Durée** | {duration_str} |
| **Catégorie** | {category} |
| **Mode d'analyse** | {mode} |
| **Date d'analyse** | {date_str} |

"""

    if video_url:
        content += f"🔗 [Voir la vidéo]({video_url})\n\n"
    
    if thumbnail_url:
        content += f"![Thumbnail]({thumbnail_url})\n\n"
    
    content += "---\n\n## 📋 Synthèse\n\n"
    content += summary + "\n\n"
    
    # Score de fiabilité
    if reliability_score is not None:
        emoji = "✅" if reliability_score >= 70 else "⚖️" if reliability_score >= 50 else "⚠️"
        content += f"---\n\n## 📊 Score de fiabilité\n\n{emoji} **{reliability_score}/100**\n\n"
    
    # Entités extraites
    if entities:
        content += "---\n\n## 🏷️ Entités extraites\n\n"
        
        if entities.get("concepts"):
            content += "### 💡 Concepts clés\n"
            for concept in entities["concepts"][:10]:
                content += f"- {concept}\n"
            content += "\n"
        
        if entities.get("persons"):
            content += "### 👤 Personnes mentionnées\n"
            for person in entities["persons"][:10]:
                content += f"- {person}\n"
            content += "\n"
        
        if entities.get("organizations"):
            content += "### 🏢 Organisations\n"
            for org in entities["organizations"][:10]:
                content += f"- {org}\n"
            content += "\n"
    
    # Flashcards
    if flashcards:
        content += "---\n\n## 📚 Flashcards de révision\n\n"
        for i, card in enumerate(flashcards[:10], 1):
            content += f"### Carte {i}\n"
            content += f"**Q:** {card.get('front', card.get('question', ''))}\n\n"
            content += f"**R:** {card.get('back', card.get('answer', ''))}\n\n"
    
    content += """---

*Généré par [Deep Sight](https://deepsightsynthesis.com) — Analyse intelligente de vidéos YouTube*
"""
    
    return content


# ═══════════════════════════════════════════════════════════════════════════════
# 📄 EXPORT DOCX
# ═══════════════════════════════════════════════════════════════════════════════

def export_to_docx(
    title: str,
    channel: str,
    category: str,
    mode: str,
    summary: str,
    video_url: str = "",
    duration: int = 0,
    entities: Dict = None,
    reliability_score: float = None,
    created_at: datetime = None,
    flashcards: List[Dict] = None
) -> Optional[bytes]:
    """Exporte l'analyse en format DOCX"""

    if not DOCX_AVAILABLE:
        return None

    # Imports lazy — chargés uniquement quand DOCX_AVAILABLE=True
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()
    
    # Titre principal
    title_para = doc.add_heading("🤿 Deep Sight — Analyse", 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sous-titre
    subtitle = doc.add_paragraph()
    subtitle.add_run("Analyse intelligente de vidéos YouTube").italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Section vidéo
    doc.add_heading("📺 Vidéo analysée", level=1)
    
    duration_str = format_duration(duration)
    date_str = created_at.strftime("%d/%m/%Y à %H:%M") if created_at else datetime.now().strftime("%d/%m/%Y à %H:%M")
    
    # Tableau d'infos
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    info = [
        ("Titre", title),
        ("Chaîne", channel),
        ("Durée", duration_str),
        ("Catégorie", category),
        ("Mode", mode),
        ("Analysé le", date_str)
    ]
    
    for i, (label, value) in enumerate(info):
        cells = table.rows[i].cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = value
    
    if video_url:
        p = doc.add_paragraph()
        p.add_run("🔗 URL: ").bold = True
        p.add_run(video_url)
    
    doc.add_paragraph()
    
    # Section synthèse
    doc.add_heading("📋 Synthèse", level=1)
    
    # Nettoyer et ajouter le contenu
    for paragraph in summary.split('\n\n'):
        if paragraph.strip():
            # Détecter les headers markdown
            if paragraph.startswith('## '):
                doc.add_heading(paragraph[3:].strip(), level=2)
            elif paragraph.startswith('### '):
                doc.add_heading(paragraph[4:].strip(), level=3)
            else:
                p = doc.add_paragraph(paragraph.strip())
    
    # Score de fiabilité
    if reliability_score is not None:
        doc.add_paragraph()
        doc.add_heading("📊 Score de fiabilité", level=1)
        emoji = "✅" if reliability_score >= 70 else "⚖️" if reliability_score >= 50 else "⚠️"
        p = doc.add_paragraph()
        p.add_run(f"{emoji} {reliability_score}/100").bold = True
    
    # Entités
    if entities:
        doc.add_paragraph()
        doc.add_heading("🏷️ Entités extraites", level=1)
        
        if entities.get("concepts"):
            doc.add_heading("Concepts clés", level=2)
            for concept in entities["concepts"][:10]:
                doc.add_paragraph(concept, style='List Bullet')
        
        if entities.get("persons"):
            doc.add_heading("Personnes", level=2)
            for person in entities["persons"][:10]:
                doc.add_paragraph(person, style='List Bullet')
    
    # Flashcards
    if flashcards:
        doc.add_paragraph()
        doc.add_heading("📚 Flashcards de révision", level=1)
        for i, card in enumerate(flashcards[:10], 1):
            doc.add_heading(f"Carte {i}", level=2)
            p = doc.add_paragraph()
            p.add_run("Question: ").bold = True
            p.add_run(card.get('front', card.get('question', '')))
            p = doc.add_paragraph()
            p.add_run("Réponse: ").bold = True
            p.add_run(card.get('back', card.get('answer', '')))
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run("Généré par Deep Sight — deepsightsynthesis.com").italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sauvegarder en bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# 📄 EXPORT PDF (ReportLab Fallback)
# ═══════════════════════════════════════════════════════════════════════════════

def export_to_pdf_reportlab(
    title: str,
    channel: str,
    category: str,
    mode: str,
    summary: str,
    video_url: str = "",
    duration: int = 0,
    entities: Dict = None,
    reliability_score: float = None,
    created_at: datetime = None
) -> Optional[bytes]:
    """Export PDF de fallback avec ReportLab (moins stylé)"""

    if not REPORTLAB_AVAILABLE:
        return None

    # Imports lazy — chargés uniquement quand REPORTLAB_AVAILABLE=True
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.units import cm

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )
    
    # Styles
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=HexColor(COLORS["primary"]),
        spaceAfter=20,
        alignment=1
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=HexColor(COLORS["primary"]),
        spaceBefore=15,
        spaceAfter=10
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        spaceAfter=8
    )
    
    info_style = ParagraphStyle(
        'InfoStyle',
        parent=styles['Normal'],
        fontSize=10,
        textColor=HexColor("#666666")
    )
    
    # Contenu
    story = []
    
    # Titre
    story.append(Paragraph("🤿 Deep Sight — Analyse", title_style))
    story.append(Spacer(1, 20))
    
    # Infos vidéo
    story.append(Paragraph("📺 Vidéo analysée", heading_style))
    
    duration_str = format_duration(duration)
    date_str = created_at.strftime("%d/%m/%Y à %H:%M") if created_at else datetime.now().strftime("%d/%m/%Y à %H:%M")
    
    info_text = f"""
    <b>Titre:</b> {title}<br/>
    <b>Chaîne:</b> {channel}<br/>
    <b>Durée:</b> {duration_str}<br/>
    <b>Catégorie:</b> {category}<br/>
    <b>Mode:</b> {mode}<br/>
    <b>Analysé le:</b> {date_str}
    """
    story.append(Paragraph(info_text, info_style))
    
    if video_url:
        story.append(Paragraph(f"<b>URL:</b> {video_url}", info_style))
    
    story.append(Spacer(1, 20))
    
    # Synthèse
    story.append(Paragraph("📋 Synthèse", heading_style))
    
    # Nettoyer le markdown pour PDF
    summary_clean = summary.replace('**', '')
    summary_clean = re.sub(r'^##+ ', '', summary_clean, flags=re.MULTILINE)
    
    for paragraph in summary_clean.split('\n\n'):
        if paragraph.strip():
            safe_text = paragraph.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(safe_text, body_style))
    
    # Score de fiabilité
    if reliability_score is not None:
        story.append(Spacer(1, 15))
        story.append(Paragraph("📊 Score de fiabilité", heading_style))
        emoji = "✅" if reliability_score >= 70 else "⚖️" if reliability_score >= 50 else "⚠️"
        story.append(Paragraph(f"<b>{emoji} {reliability_score}/100</b>", body_style))
    
    # Footer
    story.append(Spacer(1, 30))
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=8,
        textColor=HexColor("#999999"),
        alignment=1
    )
    story.append(Paragraph("Généré par Deep Sight — deepsightsynthesis.com", footer_style))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# 📄 EXPORT PDF (Main - tries WeasyPrint first, then ReportLab)
# ═══════════════════════════════════════════════════════════════════════════════

def export_to_pdf(
    title: str,
    channel: str,
    category: str,
    mode: str,
    summary: str,
    video_url: str = "",
    duration: int = 0,
    thumbnail_url: str = "",
    entities: Dict = None,
    reliability_score: float = None,
    created_at: datetime = None,
    flashcards: List[Dict] = None,
    sources: List[Dict] = None,
    export_type: str = "full"
) -> Optional[bytes]:
    """
    Exporte l'analyse en format PDF.
    Utilise WeasyPrint si disponible, sinon ReportLab en fallback.
    
    Args:
        export_type: "full" | "summary" | "flashcards" | "study"
    """
    
    # Try WeasyPrint first (beautiful HTML→PDF)
    if weasyprint_available():
        pdf = generate_pdf_weasyprint(
            title=title,
            channel=channel,
            category=category,
            mode=mode,
            summary=summary,
            video_url=video_url,
            duration=duration,
            thumbnail_url=thumbnail_url,
            entities=entities,
            reliability_score=reliability_score,
            created_at=created_at,
            flashcards=flashcards,
            sources=sources,
            export_type=export_type
        )
        if pdf:
            return pdf
        print("⚠️ WeasyPrint failed, falling back to ReportLab", flush=True)
    
    # Fallback to ReportLab
    if REPORTLAB_AVAILABLE:
        return export_to_pdf_reportlab(
            title=title,
            channel=channel,
            category=category,
            mode=mode,
            summary=summary,
            video_url=video_url,
            duration=duration,
            entities=entities,
            reliability_score=reliability_score,
            created_at=created_at
        )
    
    return None


# ═══════════════════════════════════════════════════════════════════════════════
# 📊 EXPORT CSV
# ═══════════════════════════════════════════════════════════════════════════════

def export_to_csv(
    title: str,
    channel: str,
    category: str,
    mode: str,
    summary: str,
    video_url: str = "",
    duration: int = 0,
    entities: Dict = None,
    reliability_score: float = None,
    created_at: datetime = None
) -> str:
    """Exporte l'analyse en format CSV (structured data)"""

    # Formatage durée
    if duration:
        hours, remainder = divmod(duration, 3600)
        minutes, seconds = divmod(remainder, 60)
        duration_str = f"{hours}h{minutes:02d}m" if hours else f"{minutes}m{seconds:02d}s"
    else:
        duration_str = "N/A"

    date_str = created_at.strftime("%Y-%m-%d %H:%M:%S") if created_at else datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    buffer = io.StringIO()
    writer = csv.writer(buffer, quoting=csv.QUOTE_ALL)

    # Header
    writer.writerow(["Deep Sight - Export d'analyse"])
    writer.writerow([])

    # Métadonnées
    writer.writerow(["Propriété", "Valeur"])
    writer.writerow(["Titre", title])
    writer.writerow(["Chaîne", channel])
    writer.writerow(["Durée", duration_str])
    writer.writerow(["Catégorie", category])
    writer.writerow(["Mode d'analyse", mode])
    writer.writerow(["URL", video_url])
    writer.writerow(["Date d'analyse", date_str])

    if reliability_score is not None:
        writer.writerow(["Score de fiabilité", f"{reliability_score}/100"])

    writer.writerow([])

    # Synthèse (nettoyée du markdown)
    summary_clean = re.sub(r'^##+ ', '', summary, flags=re.MULTILINE)
    summary_clean = summary_clean.replace('**', '').replace('*', '')
    writer.writerow(["Synthèse"])
    writer.writerow([summary_clean])

    writer.writerow([])

    # Entités
    if entities:
        if entities.get("concepts"):
            writer.writerow(["Concepts clés"])
            for concept in entities["concepts"][:15]:
                writer.writerow([concept])
            writer.writerow([])

        if entities.get("persons"):
            writer.writerow(["Personnes mentionnées"])
            for person in entities["persons"][:15]:
                writer.writerow([person])
            writer.writerow([])

        if entities.get("organizations"):
            writer.writerow(["Organisations"])
            for org in entities["organizations"][:15]:
                writer.writerow([org])

    writer.writerow([])
    writer.writerow(["Généré par Deep Sight - deepsightsynthesis.com"])

    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# 📊 EXPORT EXCEL (.xlsx)
# ═══════════════════════════════════════════════════════════════════════════════

def export_to_excel(
    title: str,
    channel: str,
    category: str,
    mode: str,
    summary: str,
    video_url: str = "",
    duration: int = 0,
    entities: Dict = None,
    reliability_score: float = None,
    created_at: datetime = None
) -> Optional[bytes]:
    """Exporte l'analyse en format Excel (.xlsx)"""

    if not EXCEL_AVAILABLE:
        return None

    # Imports lazy — chargés uniquement quand EXCEL_AVAILABLE=True
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Analyse Deep Sight"

    # Styles
    header_font = Font(name='Arial', size=16, bold=True, color="0D4F4F")
    subheader_font = Font(name='Arial', size=12, bold=True, color="0D4F4F")
    label_font = Font(name='Arial', size=10, bold=True)
    value_font = Font(name='Arial', size=10)
    header_fill = PatternFill(start_color="E8F4F4", end_color="E8F4F4", fill_type="solid")

    thin_border = Border(
        left=Side(style='thin', color='CCCCCC'),
        right=Side(style='thin', color='CCCCCC'),
        top=Side(style='thin', color='CCCCCC'),
        bottom=Side(style='thin', color='CCCCCC')
    )

    # Formatage durée
    if duration:
        hours, remainder = divmod(duration, 3600)
        minutes, seconds = divmod(remainder, 60)
        duration_str = f"{hours}h{minutes:02d}m" if hours else f"{minutes}m{seconds:02d}s"
    else:
        duration_str = "N/A"

    date_str = created_at.strftime("%d/%m/%Y %H:%M") if created_at else datetime.now().strftime("%d/%m/%Y %H:%M")

    row = 1

    # Titre principal
    ws.merge_cells('A1:D1')
    cell = ws['A1']
    cell.value = "🤿 Deep Sight — Analyse"
    cell.font = header_font
    cell.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 30
    row = 3

    # Section Vidéo
    ws.merge_cells(f'A{row}:D{row}')
    cell = ws[f'A{row}']
    cell.value = "📺 Vidéo analysée"
    cell.font = subheader_font
    cell.fill = header_fill
    row += 1

    # Données de la vidéo
    video_data = [
        ("Titre", title),
        ("Chaîne", channel),
        ("Durée", duration_str),
        ("Catégorie", category),
        ("Mode d'analyse", mode),
        ("URL", video_url),
        ("Date d'analyse", date_str),
    ]

    if reliability_score is not None:
        emoji = "✅" if reliability_score >= 70 else "⚖️" if reliability_score >= 50 else "⚠️"
        video_data.append(("Score de fiabilité", f"{emoji} {reliability_score}/100"))

    for label, value in video_data:
        ws[f'A{row}'] = label
        ws[f'A{row}'].font = label_font
        ws[f'A{row}'].border = thin_border
        ws.merge_cells(f'B{row}:D{row}')
        ws[f'B{row}'] = value
        ws[f'B{row}'].font = value_font
        ws[f'B{row}'].border = thin_border
        row += 1

    row += 1

    # Section Synthèse
    ws.merge_cells(f'A{row}:D{row}')
    cell = ws[f'A{row}']
    cell.value = "📋 Synthèse"
    cell.font = subheader_font
    cell.fill = header_fill
    row += 1

    # Contenu de la synthèse (nettoyé)
    summary_clean = re.sub(r'^##+ ', '', summary, flags=re.MULTILINE)
    summary_clean = summary_clean.replace('**', '').replace('*', '')

    ws.merge_cells(f'A{row}:D{row}')
    cell = ws[f'A{row}']
    cell.value = summary_clean[:32000]  # Excel cell limit
    cell.font = value_font
    cell.alignment = Alignment(wrap_text=True, vertical='top')
    ws.row_dimensions[row].height = min(400, max(50, len(summary_clean) // 5))
    row += 2

    # Section Entités
    if entities:
        ws.merge_cells(f'A{row}:D{row}')
        cell = ws[f'A{row}']
        cell.value = "🏷️ Entités extraites"
        cell.font = subheader_font
        cell.fill = header_fill
        row += 1

        col_offset = 0
        if entities.get("concepts"):
            ws[f'A{row}'] = "Concepts clés"
            ws[f'A{row}'].font = label_font
            for i, concept in enumerate(entities["concepts"][:15]):
                ws[f'A{row + 1 + i}'] = concept
                ws[f'A{row + 1 + i}'].font = value_font

        if entities.get("persons"):
            ws[f'B{row}'] = "Personnes"
            ws[f'B{row}'].font = label_font
            for i, person in enumerate(entities["persons"][:15]):
                ws[f'B{row + 1 + i}'] = person
                ws[f'B{row + 1 + i}'].font = value_font

        if entities.get("organizations"):
            ws[f'C{row}'] = "Organisations"
            ws[f'C{row}'].font = label_font
            for i, org in enumerate(entities["organizations"][:15]):
                ws[f'C{row + 1 + i}'] = org
                ws[f'C{row + 1 + i}'].font = value_font

        row += 17

    # Footer
    ws.merge_cells(f'A{row}:D{row}')
    cell = ws[f'A{row}']
    cell.value = "Généré par Deep Sight — deepsightsynthesis.com"
    cell.font = Font(name='Arial', size=8, italic=True, color='999999')
    cell.alignment = Alignment(horizontal='center')

    # Ajuster largeur des colonnes
    ws.column_dimensions['A'].width = 20
    ws.column_dimensions['B'].width = 30
    ws.column_dimensions['C'].width = 25
    ws.column_dimensions['D'].width = 25

    # Sauvegarder en bytes
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()



# ═══════════════════════════════════════════════════════════════════════════════
# 🔧 FONCTION PRINCIPALE D'EXPORT
# ═══════════════════════════════════════════════════════════════════════════════

def export_summary(
    format: str,
    title: str,
    channel: str,
    category: str,
    mode: str,
    summary: str,
    video_url: str = "",
    duration: int = 0,
    thumbnail_url: str = "",
    entities: Dict = None,
    reliability_score: float = None,
    created_at: datetime = None,
    flashcards: List[Dict] = None,
    sources: List[Dict] = None,
    pdf_export_type: str = "full"
) -> Tuple[Optional[bytes | str], str, str]:
    """
    Exporte un résumé dans le format demandé.
    
    Args:
        format: txt, md, docx, pdf
        pdf_export_type: full, summary, flashcards, study (pour PDF uniquement)
    
    Returns:
        Tuple (content, filename, mimetype)
    """
    
    # Générer un nom de fichier safe
    timestamp = datetime.now().strftime("%Y%m%d")
    base_filename = clean_filename(title, timestamp)
    
    if format == "txt":
        content = export_to_txt(
            title, channel, category, mode, summary,
            video_url, duration, created_at
        )
        return content, f"{base_filename}.txt", "text/plain"
    
    elif format == "md":
        content = export_to_markdown(
            title, channel, category, mode, summary,
            video_url, duration, thumbnail_url, entities,
            reliability_score, created_at, flashcards
        )
        return content, f"{base_filename}.md", "text/markdown"
    
    elif format == "docx":
        if not DOCX_AVAILABLE:
            return None, "", ""
        content = export_to_docx(
            title, channel, category, mode, summary,
            video_url, duration, entities, reliability_score, 
            created_at, flashcards
        )
        return content, f"{base_filename}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    elif format == "pdf":
        content = export_to_pdf(
            title=title,
            channel=channel,
            category=category,
            mode=mode,
            summary=summary,
            video_url=video_url,
            duration=duration,
            thumbnail_url=thumbnail_url,
            entities=entities,
            reliability_score=reliability_score,
            created_at=created_at,
            flashcards=flashcards,
            sources=sources,
            export_type=pdf_export_type
        )
        if content is None:
            return None, "", ""
        
        # Ajouter le type dans le nom de fichier
        type_suffix = "" if pdf_export_type == "full" else f"_{pdf_export_type}"
        return content, f"{base_filename}{type_suffix}.pdf", "application/pdf"
        return content, f"deepsight_{safe_title}_{timestamp}.pdf", "application/pdf"

    elif format == "csv":
        content = export_to_csv(
            title, channel, category, mode, summary,
            video_url, duration, entities, reliability_score, created_at
        )
        return content, f"deepsight_{safe_title}_{timestamp}.csv", "text/csv"

    elif format == "xlsx":
        if not EXCEL_AVAILABLE:
            return None, "", ""
        content = export_to_excel(
            title, channel, category, mode, summary,
            video_url, duration, entities, reliability_score, created_at
        )
        return content, f"deepsight_{safe_title}_{timestamp}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    else:
        return None, "", ""


def get_available_formats() -> List[str]:
    """Retourne la liste des formats d'export disponibles"""
    formats = ["txt", "md", "csv"]  # CSV is always available (stdlib)
    if DOCX_AVAILABLE:
        formats.append("docx")
    if weasyprint_available() or REPORTLAB_AVAILABLE:
        formats.append("pdf")
    if EXCEL_AVAILABLE:
        formats.append("xlsx")
    if is_audio_export_available():
        formats.append("audio")
    return formats


def get_pdf_export_options() -> List[Dict]:
    """Retourne les options d'export PDF disponibles (lazy)"""
    if _ensure_pdf():
        return _pdf_module.PDF_EXPORT_OPTIONS
    return []


# ═══════════════════════════════════════════════════════════════════════════════
# 🔊 AUDIO EXPORT (ElevenLabs TTS)
# ═══════════════════════════════════════════════════════════════════════════════

import uuid
import logging
import glob as glob_module
import time as _time

_audio_logger = logging.getLogger(__name__)

AUDIO_TMP_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "tmp", "audio")
os.makedirs(AUDIO_TMP_DIR, exist_ok=True)

# Cleanup threshold: 24 hours
_AUDIO_MAX_AGE_SECONDS = 86400


def is_audio_export_available() -> bool:
    """Check if audio export is available (ElevenLabs key configured)."""
    try:
        from core.config import get_elevenlabs_key
        return bool(get_elevenlabs_key())
    except ImportError:
        return False


def cleanup_old_audio_files() -> int:
    """Remove audio files older than 24h. Returns count of removed files."""
    removed = 0
    now = _time.time()
    for filepath in glob_module.glob(os.path.join(AUDIO_TMP_DIR, "*.mp3")):
        try:
            if now - os.path.getmtime(filepath) > _AUDIO_MAX_AGE_SECONDS:
                os.remove(filepath)
                removed += 1
        except OSError:
            pass
    return removed


def build_narrative_text(
    title: str,
    channel: str,
    summary: str,
    mode: str = "",
    condensed: bool = False,
) -> str:
    """
    Build a natural-sounding narrative from the analysis content.
    Designed to be read aloud by TTS.

    If condensed=True, truncate to ~300 words (~2 min at 150 wpm).
    """
    try:
        from tts.service import clean_text_for_tts
    except ImportError:
        clean_text_for_tts = lambda t, **kw: t

    # Build intro
    parts = []
    parts.append(f"Voici l'analyse de la vidéo intitulée {title}")
    if channel:
        parts.append(f"publiée par {channel}.")
    else:
        parts.append(".")

    # Clean the main summary content
    cleaned_summary = clean_text_for_tts(summary, strip_questions=True)

    # Condensed mode: keep ~300 words (≈2 min audio)
    if condensed and cleaned_summary:
        words = cleaned_summary.split()
        if len(words) > 300:
            # Cut at last paragraph break before 300 words
            truncated = " ".join(words[:300])
            last_period = truncated.rfind(".")
            if last_period > len(truncated) * 0.5:
                truncated = truncated[:last_period + 1]
            cleaned_summary = truncated

    if cleaned_summary:
        parts.append(cleaned_summary)

    # Build outro
    parts.append("Voilà qui conclut cette analyse. Document généré par DeepSight.")

    full_text = " ".join(parts)

    # Remove any remaining markdown artifacts
    full_text = re.sub(r'#{1,6}\s+', '', full_text)
    full_text = re.sub(r'\*\*(.+?)\*\*', r'\1', full_text)
    full_text = re.sub(r'\*(.+?)\*', r'\1', full_text)
    full_text = re.sub(r'`[^`]*`', '', full_text)
    full_text = re.sub(r'---+', '', full_text)
    full_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', full_text)  # [text](url) → text
    full_text = re.sub(r'\s{2,}', ' ', full_text)

    return full_text.strip()


def chunk_text_for_tts(text: str, max_chunk_size: int = 4500) -> List[str]:
    """
    Split text into chunks at sentence boundaries.
    ElevenLabs has a ~5000 char limit per request.
    """
    if len(text) <= max_chunk_size:
        return [text]

    chunks = []
    remaining = text

    while remaining:
        if len(remaining) <= max_chunk_size:
            chunks.append(remaining)
            break

        # Find last sentence boundary within limit
        cut_point = -1
        for sep in ['. ', '.\n', '! ', '? ', ';\n']:
            idx = remaining.rfind(sep, 0, max_chunk_size)
            if idx > cut_point:
                cut_point = idx + len(sep)

        if cut_point <= 0:
            # No sentence boundary found, cut at space
            cut_point = remaining.rfind(' ', 0, max_chunk_size)
            if cut_point <= 0:
                cut_point = max_chunk_size

        chunks.append(remaining[:cut_point].strip())
        remaining = remaining[cut_point:].strip()

    return [c for c in chunks if c]


async def export_to_audio(
    title: str,
    channel: str,
    summary: str,
    mode: str = "",
    voice_id: str = "",
    speed: float = 1.0,
    condensed: bool = False,
) -> Optional[Dict[str, Any]]:
    """
    Generate MP3 audio from analysis content via ElevenLabs TTS.

    Args:
        condensed: If True, truncate to ~300 words (~2 min audio).

    Returns:
        dict with { file_id, file_path, duration_estimate } or None on failure.
    """
    from core.config import get_elevenlabs_key

    api_key = get_elevenlabs_key()
    if not api_key:
        _audio_logger.error("ElevenLabs API key not configured for audio export")
        return None

    # Cleanup old files opportunistically
    cleanup_old_audio_files()

    # Build narrative text
    narrative = build_narrative_text(title, channel, summary, mode, condensed=condensed)
    if not narrative or len(narrative) < 10:
        _audio_logger.error("Narrative text too short for audio export")
        return None

    # Determine voice
    if not voice_id:
        try:
            from tts.service import get_voice_id, DEFAULT_MODEL_ID
            voice_id = get_voice_id("fr", "female")
            model_id = DEFAULT_MODEL_ID
        except ImportError:
            voice_id = "pFZP5JQG7iQjIQuC4Bku"  # Lily default
            model_id = "eleven_multilingual_v2"
    else:
        model_id = "eleven_multilingual_v2"

    # Chunk the text
    chunks = chunk_text_for_tts(narrative)
    _audio_logger.info(f"Audio export: {len(narrative)} chars → {len(chunks)} chunks")

    # Generate audio for each chunk
    file_id = str(uuid.uuid4())
    file_path = os.path.join(AUDIO_TMP_DIR, f"{file_id}.mp3")

    try:
        async with shared_http_client() as client:
            audio_parts: list[bytes] = []

            for i, chunk_text in enumerate(chunks):
                url = f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}"
                payload = {
                    "text": chunk_text,
                    "model_id": model_id,
                    "voice_settings": {
                        "stability": 0.5,
                        "similarity_boost": 0.75,
                        "style": 0.3,
                        "use_speaker_boost": True,
                        "speed": speed,
                    },
                }
                headers = {
                    "xi-api-key": api_key,
                    "Content-Type": "application/json",
                    "Accept": "audio/mpeg",
                }

                response = await client.post(url, json=payload, headers=headers, timeout=120.0)

                if response.status_code != 200:
                    _audio_logger.error(
                        f"ElevenLabs API error on chunk {i+1}/{len(chunks)}: "
                        f"{response.status_code} {response.text[:200]}"
                    )
                    return None

                audio_parts.append(response.content)

            # Concatenate MP3 chunks (simple byte concatenation works for MP3)
            with open(file_path, "wb") as f:
                for part in audio_parts:
                    f.write(part)

        # Estimate duration: ~150 words/min for French TTS
        word_count = len(narrative.split())
        duration_estimate = int((word_count / 150) * 60)

        return {
            "file_id": file_id,
            "file_path": file_path,
            "duration_estimate": duration_estimate,
        }

    except httpx.TimeoutException:
        _audio_logger.error("ElevenLabs API timeout during audio export")
        return None
    except Exception as e:
        _audio_logger.error(f"Audio export error: {e}")
        # Cleanup partial file
        if os.path.exists(file_path):
            os.remove(file_path)
        return None


def get_audio_file_path(file_id: str) -> Optional[str]:
    """Get the path to a temporary audio file, validating the file_id format."""
    # Validate UUID format to prevent path traversal
    try:
        uuid.UUID(file_id)
    except ValueError:
        return None

    file_path = os.path.join(AUDIO_TMP_DIR, f"{file_id}.mp3")
    if os.path.exists(file_path):
        return file_path
    return None
