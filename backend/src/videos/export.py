"""
╔════════════════════════════════════════════════════════════════════════════════════╗
║  📤 EXPORT SERVICE — Export des résumés en PDF, DOCX, MD, TXT                      ║
╚════════════════════════════════════════════════════════════════════════════════════╝
"""

import os
import io
import re
from datetime import datetime
from typing import Optional, Dict, Any

# ═══════════════════════════════════════════════════════════════════════════════
# 📝 MARKDOWN EXPORT
# ═══════════════════════════════════════════════════════════════════════════════

def export_to_markdown(
    title: str,
    channel: str,
    summary: str,
    video_url: str,
    category: str = "",
    mode: str = "",
    lang: str = "fr",
    created_at: datetime = None
) -> str:
    """Exporte un résumé en Markdown"""
    
    date_str = created_at.strftime("%d/%m/%Y %H:%M") if created_at else datetime.now().strftime("%d/%m/%Y %H:%M")
    
    header_fr = f"""# 🤿 Deep Sight — Synthèse

## 📺 {title}

**Chaîne:** {channel}  
**URL:** {video_url}  
**Catégorie:** {category}  
**Mode:** {mode}  
**Date d'analyse:** {date_str}

---

"""
    
    header_en = f"""# 🤿 Deep Sight — Summary

## 📺 {title}

**Channel:** {channel}  
**URL:** {video_url}  
**Category:** {category}  
**Mode:** {mode}  
**Analysis date:** {date_str}

---

"""
    
    header = header_fr if lang == "fr" else header_en
    
    footer = """

---

*Généré par Deep Sight — Analyse intelligente de vidéos YouTube*
"""
    
    return header + summary + footer


def export_to_txt(
    title: str,
    channel: str,
    summary: str,
    video_url: str,
    **kwargs
) -> str:
    """Exporte un résumé en texte brut"""
    
    # Nettoyer le markdown
    clean_text = summary
    clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', clean_text)  # **bold**
    clean_text = re.sub(r'\*([^*]+)\*', r'\1', clean_text)  # *italic*
    clean_text = re.sub(r'#{1,6}\s*', '', clean_text)  # headers
    clean_text = re.sub(r'^\s*[-*]\s+', '• ', clean_text, flags=re.MULTILINE)  # bullets
    
    return f"""DEEP SIGHT — SYNTHÈSE
{'=' * 50}

Titre: {title}
Chaîne: {channel}
URL: {video_url}

{'=' * 50}

{clean_text}

{'=' * 50}
Généré par Deep Sight
"""


# ═══════════════════════════════════════════════════════════════════════════════
# 📄 DOCX EXPORT
# ═══════════════════════════════════════════════════════════════════════════════

def export_to_docx(
    title: str,
    channel: str,
    summary: str,
    video_url: str,
    category: str = "",
    mode: str = "",
    lang: str = "fr",
    created_at: datetime = None
) -> io.BytesIO:
    """Exporte un résumé en DOCX"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx is required for DOCX export")
    
    doc = Document()
    
    # Style du titre
    title_para = doc.add_heading("🤿 Deep Sight — Synthèse", level=0)
    
    # Infos vidéo
    doc.add_heading(f"📺 {title}", level=1)
    
    info_para = doc.add_paragraph()
    info_para.add_run("Chaîne: ").bold = True
    info_para.add_run(f"{channel}\n")
    info_para.add_run("URL: ").bold = True
    info_para.add_run(f"{video_url}\n")
    info_para.add_run("Catégorie: ").bold = True
    info_para.add_run(f"{category}\n")
    info_para.add_run("Mode: ").bold = True
    info_para.add_run(f"{mode}\n")
    
    # Ligne de séparation
    doc.add_paragraph("─" * 50)
    
    # Contenu
    doc.add_heading("Synthèse", level=2)
    
    # Parser le markdown basique
    lines = summary.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        # Headers
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        # Bullets
        elif line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph(line[2:], style='List Bullet')
        # Normal text
        else:
            # Gérer le bold
            para = doc.add_paragraph()
            parts = re.split(r'\*\*([^*]+)\*\*', line)
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Bold
                    para.add_run(part).bold = True
                else:
                    para.add_run(part)
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph("─" * 50)
    footer = doc.add_paragraph("Généré par Deep Sight — Analyse intelligente de vidéos YouTube")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sauvegarder dans un buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


# ═══════════════════════════════════════════════════════════════════════════════
# 📕 PDF EXPORT
# ═══════════════════════════════════════════════════════════════════════════════

def export_to_pdf(
    title: str,
    channel: str,
    summary: str,
    video_url: str,
    category: str = "",
    mode: str = "",
    lang: str = "fr",
    created_at: datetime = None
) -> io.BytesIO:
    """Exporte un résumé en PDF avec ReportLab"""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch, cm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, 
            HRFlowable, ListFlowable, ListItem
        )
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    except ImportError:
        raise ImportError("reportlab is required for PDF export")
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )
    
    # Styles
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=12,
        textColor=colors.HexColor('#0d6e6e'),  # Teal Deep Sight
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=8,
        textColor=colors.HexColor('#0d6e6e')
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6,
        alignment=TA_JUSTIFY
    )
    
    info_style = ParagraphStyle(
        'Info',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.grey
    )
    
    # Contenu
    story = []
    
    # Titre principal
    story.append(Paragraph("🤿 Deep Sight — Synthèse", title_style))
    story.append(Spacer(1, 12))
    
    # Infos vidéo
    story.append(Paragraph(f"📺 <b>{title}</b>", heading_style))
    story.append(Paragraph(f"<b>Chaîne:</b> {channel}", info_style))
    story.append(Paragraph(f"<b>URL:</b> {video_url}", info_style))
    story.append(Paragraph(f"<b>Catégorie:</b> {category} | <b>Mode:</b> {mode}", info_style))
    story.append(Spacer(1, 12))
    
    # Ligne de séparation
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#0d6e6e')))
    story.append(Spacer(1, 12))
    
    # Parser le contenu markdown
    lines = summary.split('\n')
    current_list_items = []
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Vider la liste en cours si existe
            if current_list_items:
                story.append(ListFlowable(current_list_items, bulletType='bullet'))
                current_list_items = []
            story.append(Spacer(1, 6))
            continue
        
        # Nettoyer le markdown pour PDF
        line = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', line)  # Bold
        line = re.sub(r'\*([^*]+)\*', r'<i>\1</i>', line)  # Italic
        line = re.sub(r'`([^`]+)`', r'<font face="Courier">\1</font>', line)  # Code
        
        # Headers
        if line.startswith('### '):
            if current_list_items:
                story.append(ListFlowable(current_list_items, bulletType='bullet'))
                current_list_items = []
            story.append(Paragraph(line[4:], heading_style))
        elif line.startswith('## '):
            if current_list_items:
                story.append(ListFlowable(current_list_items, bulletType='bullet'))
                current_list_items = []
            story.append(Paragraph(line[3:], heading_style))
        elif line.startswith('# '):
            if current_list_items:
                story.append(ListFlowable(current_list_items, bulletType='bullet'))
                current_list_items = []
            story.append(Paragraph(line[2:], title_style))
        # Bullets
        elif line.startswith('- ') or line.startswith('* '):
            current_list_items.append(ListItem(Paragraph(line[2:], body_style)))
        # Normal
        else:
            if current_list_items:
                story.append(ListFlowable(current_list_items, bulletType='bullet'))
                current_list_items = []
            story.append(Paragraph(line, body_style))
    
    # Vider la liste finale si existe
    if current_list_items:
        story.append(ListFlowable(current_list_items, bulletType='bullet'))
    
    # Footer
    story.append(Spacer(1, 24))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
    story.append(Spacer(1, 6))
    
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.grey,
        alignment=TA_CENTER
    )
    story.append(Paragraph("Généré par Deep Sight — Analyse intelligente de vidéos YouTube", footer_style))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    
    return buffer


# ═══════════════════════════════════════════════════════════════════════════════
# 🎯 EXPORT PRINCIPAL
# ═══════════════════════════════════════════════════════════════════════════════

def export_summary(
    format: str,
    title: str,
    channel: str,
    summary: str,
    video_url: str,
    category: str = "",
    mode: str = "",
    lang: str = "fr",
    created_at: datetime = None
) -> tuple:
    """
    Exporte un résumé dans le format demandé.
    Retourne: (content, content_type, extension)
    """
    
    params = {
        "title": title,
        "channel": channel,
        "summary": summary,
        "video_url": video_url,
        "category": category,
        "mode": mode,
        "lang": lang,
        "created_at": created_at
    }
    
    if format == "md":
        content = export_to_markdown(**params)
        return content.encode('utf-8'), "text/markdown", ".md"
    
    elif format == "txt":
        content = export_to_txt(**params)
        return content.encode('utf-8'), "text/plain", ".txt"
    
    elif format == "docx":
        buffer = export_to_docx(**params)
        return buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".docx"
    
    elif format == "pdf":
        buffer = export_to_pdf(**params)
        return buffer.getvalue(), "application/pdf", ".pdf"
    
    elif format == "json":
        import json
        data = {
            "title": title,
            "channel": channel,
            "video_url": video_url,
            "category": category,
            "mode": mode,
            "lang": lang,
            "summary": summary,
            "exported_at": datetime.now().isoformat()
        }
        return json.dumps(data, ensure_ascii=False, indent=2).encode('utf-8'), "application/json", ".json"
    
    else:
        raise ValueError(f"Unsupported export format: {format}")
